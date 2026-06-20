from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from create_chapter1_report import (
    OUTPUT_DIR,
    LOGO_PATH,
    NAVY,
    BLUE,
    DARK_BLUE,
    GRAY,
    add_body_paragraph,
    add_bullet,
    add_callout,
    add_cover,
    add_page_number,
    configure_styles,
    set_cell_margins,
    set_cell_shading,
    set_repeat_table_header,
    set_run_font,
    set_table_fixed_width,
)


OUTPUT_PATH = OUTPUT_DIR / "Bao_cao_IronZone_Chuong_1.docx"


def add_numbered_process(doc, title, steps):
    heading = doc.add_paragraph()
    heading.paragraph_format.space_before = Pt(8)
    heading.paragraph_format.space_after = Pt(5)
    run = heading.add_run(title)
    set_run_font(run, size=13, color=DARK_BLUE, bold=True)
    for step in steps:
        paragraph = doc.add_paragraph(style="List Number")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = paragraph.add_run(step)
        set_run_font(run)


def add_form_table(doc, caption_text, fields):
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    set_table_fixed_width(table, [850, 2850, 5660])
    headers = ["STT", "Thông tin thu thập", "Mô tả / dữ liệu mẫu"]
    for index, text in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, "E8EEF5")
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(text)
        set_run_font(run, size=11, color=NAVY, bold=True)
    set_repeat_table_header(table.rows[0])

    for index, (field, description) in enumerate(fields, start=1):
        cells = table.add_row().cells
        values = (str(index), field, description)
        for col, text in enumerate(values):
            paragraph = cells[col].paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if col == 0 else WD_ALIGN_PARAGRAPH.LEFT
            run = paragraph.add_run(text)
            set_run_font(run, size=11)

    caption = doc.add_paragraph(style="Figure Caption")
    caption.add_run(caption_text)


def build_document():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(1.15)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.4)
    section.footer_distance = Inches(0.4)

    configure_styles(doc)
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.paragraph_format.space_after = Pt(0)
    run = header.add_run("IRONZONE GYM | PHÂN TÍCH THIẾT KẾ HỆ THỐNG")
    set_run_font(run, size=9, color=GRAY, bold=True)
    add_page_number(section.footer.paragraphs[0])

    add_cover(doc)

    chapter = doc.add_paragraph()
    chapter.alignment = WD_ALIGN_PARAGRAPH.CENTER
    chapter.paragraph_format.space_after = Pt(6)
    run = chapter.add_run("CHƯƠNG 1")
    set_run_font(run, size=16, color=BLUE, bold=True)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(22)
    run = title.add_run("TỔNG QUAN")
    set_run_font(run, size=24, color=NAVY, bold=True)

    add_callout(
        doc,
        "Nội dung chương",
        "Giới thiệu đề tài, xác định mục tiêu và phạm vi, khảo sát cơ cấu cùng quy trình nghiệp vụ của hệ thống IronZone Gym, sau đó tổng kết để làm cơ sở cho phần phân tích ở Chương 2.",
    )

    doc.add_heading("1.1. GIỚI THIỆU", level=1)
    add_body_paragraph(
        doc,
        "Cùng với sự phát triển của công nghệ thông tin, nhu cầu số hóa hoạt động quản lý tại các cơ sở kinh doanh dịch vụ ngày càng trở nên cần thiết. Đối với phòng gym, số lượng dữ liệu phát sinh tương đối đa dạng, bao gồm thông tin hội viên, gói tập, huấn luyện viên, đơn hàng, chương trình khuyến mãi, membership, thiết bị và lịch bảo trì. Khi các dữ liệu này được lưu trữ bằng sổ sách, bảng tính hoặc nhiều công cụ riêng lẻ, nhân viên khó tra cứu thông tin đồng nhất, quản trị viên khó theo dõi tình hình kinh doanh và khách hàng phải thực hiện nhiều thao tác trực tiếp tại quầy."
    )
    add_body_paragraph(
        doc,
        "Trong thực tế, một khách hàng trước khi đăng ký thường cần tìm hiểu giá, thời hạn và quyền lợi của từng gói tập; so sánh chuyên môn của các huấn luyện viên; lựa chọn phương thức thanh toán; sau đó theo dõi thời hạn sử dụng dịch vụ. Ở chiều ngược lại, phòng gym cần ghi nhận chính xác đơn hàng, kích hoạt membership, quản lý trạng thái hội viên và tổng hợp doanh thu. Chỉ một sai lệch giữa đơn hàng và thời hạn membership cũng có thể ảnh hưởng trực tiếp đến quyền lợi khách hàng và công tác vận hành."
    )
    add_body_paragraph(
        doc,
        "Từ những vấn đề trên, đề tài “Phân tích và thiết kế hệ thống quản lý phòng gym IronZone” được lựa chọn để xây dựng một mô hình quản lý tập trung trên nền tảng web. IronZone cho phép khách truy cập xem thông tin phòng gym, gói tập và huấn luyện viên. Sau khi đăng ký tài khoản, hội viên có thể thêm gói tập vào giỏ hàng, áp dụng mã khuyến mãi, tạo đơn hàng và theo dõi membership. Quản trị viên sử dụng một khu vực riêng để quản lý hội viên, gói tập, huấn luyện viên, thiết bị, bảo trì, khuyến mãi và đơn hàng."
    )
    add_body_paragraph(
        doc,
        "Đề tài được thực hiện trong khuôn khổ môn Phân tích thiết kế hệ thống. Vì vậy, website không chỉ là sản phẩm chạy thử mà còn là mô hình nghiệp vụ dùng để minh họa các nội dung của môn học. Từ chức năng thực tế của IronZone, nhóm có thể xác định tác nhân, yêu cầu chức năng, quy tắc nghiệp vụ và quan hệ dữ liệu; sau đó xây dựng Use Case Diagram, Activity Diagram, Sequence Diagram, Class Diagram, ERD và các sơ đồ kiến trúc liên quan."
    )
    add_body_paragraph(doc, "Các lý do chính để lựa chọn đề tài gồm:")
    add_bullet(doc, "Bài toán quản lý phòng gym gần gũi, dễ khảo sát và có khả năng áp dụng trong thực tế.")
    add_bullet(doc, "Hệ thống có nhiều nhóm chức năng liên kết với nhau, phù hợp để thể hiện đầy đủ quy trình phân tích và thiết kế.")
    add_bullet(doc, "Hai nhóm tác nhân chính là hội viên và quản trị viên có quyền hạn rõ ràng, thuận lợi cho việc mô hình hóa Use Case.")
    add_bullet(doc, "Quy trình mua gói tập có nhiều bước và tác động đến nhiều đối tượng dữ liệu, phù hợp để xây dựng Activity Diagram và Sequence Diagram.")
    add_bullet(doc, "Hệ thống gồm frontend, backend và cơ sở dữ liệu, tạo điều kiện minh họa kiến trúc nhiều tầng và sự tương tác giữa các thành phần.")

    doc.add_heading("1.2. MỤC TIÊU VÀ PHẠM VI ĐỀ TÀI", level=1)
    doc.add_heading("1.2.1. Mục tiêu đề tài", level=2)
    add_body_paragraph(
        doc,
        "Mục tiêu tổng quát của đề tài là phân tích, thiết kế và cài đặt thử nghiệm một hệ thống web hỗ trợ quản lý phòng gym. Hệ thống phải mô phỏng được các nghiệp vụ tiêu biểu của IronZone, đồng thời bảo đảm các chức năng đã cài đặt có thể đối chiếu với tài liệu và sơ đồ phân tích thiết kế."
    )
    add_body_paragraph(doc, "Các mục tiêu cụ thể bao gồm:")
    add_bullet(doc, "Khảo sát và mô tả cơ cấu tổ chức, đối tượng tham gia và các quy trình nghiệp vụ chính của phòng gym.")
    add_bullet(doc, "Xác định yêu cầu chức năng, yêu cầu phi chức năng, quy tắc nghiệp vụ và phạm vi của hệ thống.")
    add_bullet(doc, "Xây dựng các sơ đồ Use Case, Activity, Sequence, Class, ERD và sơ đồ kiến trúc cho các chức năng trọng tâm.")
    add_bullet(doc, "Thiết kế cơ sở dữ liệu có quan hệ rõ ràng giữa người dùng, gói tập, đơn hàng, membership và các đối tượng vận hành.")
    add_bullet(doc, "Cài đặt giao diện khách hàng, giao diện quản trị, REST API và cơ sở dữ liệu SQLite.")
    add_bullet(doc, "Kiểm tra các luồng đăng nhập, mua gói, áp dụng khuyến mãi, tạo membership và quản trị dữ liệu.")

    doc.add_heading("1.2.2. Phạm vi đề tài", level=2)
    add_body_paragraph(doc, "Phạm vi chức năng của hệ thống được giới hạn trong các nội dung sau:")
    add_bullet(doc, "Khách truy cập: xem trang chủ, danh sách gói tập, danh sách huấn luyện viên, giới thiệu và thông tin liên hệ.")
    add_bullet(doc, "Hội viên: đăng ký, đăng nhập, cập nhật hồ sơ, quản lý giỏ hàng, mua gói tập, sử dụng mã khuyến mãi và theo dõi đơn hàng.")
    add_bullet(doc, "Membership: tự động tạo sau khi đơn hàng được xác nhận và hiển thị thời gian bắt đầu, kết thúc, trạng thái cùng huấn luyện viên liên quan nếu có.")
    add_bullet(doc, "Quản trị viên: quản lý hội viên, gói tập, huấn luyện viên, thiết bị, bảo trì, khuyến mãi và đơn hàng.")
    add_bullet(doc, "Thống kê quản trị: tổng số thành viên, membership đang hoạt động, doanh thu, đơn hàng và thiết bị cần chú ý.")

    add_body_paragraph(doc, "Giới hạn của phiên bản hiện tại:")
    add_bullet(doc, "Sản phẩm phục vụ học tập và demo, chưa được thiết kế cho hoạt động thương mại quy mô lớn.")
    add_bullet(doc, "Các phương thức thanh toán chỉ được mô phỏng ở mức lựa chọn và ghi nhận, chưa kết nối cổng thanh toán thật.")
    add_bullet(doc, "Chưa hỗ trợ đặt lịch huấn luyện viên, quản lý lớp học, chấm công nhân viên hoặc check-in bằng thiết bị phần cứng.")
    add_bullet(doc, "Cơ sở dữ liệu SQLite phù hợp cho demo nhưng còn hạn chế khi có nhiều người dùng truy cập đồng thời.")
    add_bullet(doc, "Dữ liệu mẫu được tạo bằng chương trình seed và có thể bị khởi tạo lại trong môi trường triển khai demo.")

    doc.add_heading("1.3. KHẢO SÁT HỆ THỐNG", level=1)
    doc.add_heading("1.3.1. Giới thiệu sơ lược hệ thống thông tin", level=2)
    add_body_paragraph(
        doc,
        "IronZone Gym Management System là hệ thống thông tin quản lý được xây dựng dưới dạng ứng dụng web full-stack. Mục đích của hệ thống là kết nối nhu cầu đăng ký dịch vụ của hội viên với hoạt động quản lý của phòng gym trên cùng một nền tảng. Dữ liệu được lưu trữ tập trung và được truy cập thông qua các chức năng phù hợp với quyền của từng loại tài khoản."
    )
    add_body_paragraph(
        doc,
        "Về mặt kỹ thuật, hệ thống gồm ba thành phần chính. Frontend được xây dựng bằng React và cung cấp giao diện tương tác cho người dùng. Backend sử dụng Node.js và Express để cung cấp REST API, kiểm tra dữ liệu, xử lý nghiệp vụ, xác thực và phân quyền. Cơ sở dữ liệu SQLite lưu trữ các bảng users, admins, trainers, plans, cart_items, promotions, orders, order_items, memberships, equipment, maintenance_logs, reviews và notifications."
    )
    add_body_paragraph(
        doc,
        "Hệ thống sử dụng JSON Web Token để xác thực. Khi người dùng đăng nhập thành công, client lưu token và gửi token trong tiêu đề Authorization của các yêu cầu cần bảo vệ. Server kiểm tra token, xác định tài khoản tương ứng và phân biệt người dùng thông thường với quản trị viên trước khi cho phép truy cập tài nguyên."
    )

    doc.add_heading("1.3.2. Cơ cấu tổ chức của hệ thống", level=2)
    add_body_paragraph(
        doc,
        "Trong phạm vi khảo sát, cơ cấu hoạt động của IronZone được mô hình hóa theo các bộ phận và vai trò sau:"
    )
    add_bullet(doc, "Ban quản lý / quản trị viên: theo dõi số liệu tổng quan, quản lý dữ liệu hệ thống và điều phối hoạt động.")
    add_bullet(doc, "Bộ phận tư vấn và chăm sóc hội viên: tiếp nhận thông tin đăng ký, hỗ trợ lựa chọn gói tập và xử lý yêu cầu của hội viên.")
    add_bullet(doc, "Bộ phận huấn luyện viên: cung cấp thông tin chuyên môn và thực hiện hoạt động huấn luyện theo gói dịch vụ.")
    add_bullet(doc, "Bộ phận vận hành và bảo trì: quản lý thiết bị, ghi nhận tình trạng và thực hiện các đợt kiểm tra, sửa chữa.")
    add_bullet(doc, "Hội viên / khách hàng: tìm hiểu dịch vụ, tạo tài khoản, mua gói tập và theo dõi quyền lợi.")
    add_body_paragraph(
        doc,
        "Trong phiên bản website, quản trị viên đại diện cho các bộ phận nội bộ khi thao tác dữ liệu. Việc phân tách chi tiết nhiều vai trò nhân viên chưa được cài đặt nhằm giữ phạm vi đề tài phù hợp với thời lượng học phần."
    )

    doc.add_heading("1.3.3. Quy trình nghiệp vụ bên trong hệ thống", level=2)
    add_numbered_process(
        doc,
        "a) Quy trình đăng ký và đăng nhập",
        [
            "Khách hàng truy cập trang đăng ký và nhập họ tên, email, số điện thoại cùng mật khẩu.",
            "Frontend gửi dữ liệu đăng ký đến API; server kiểm tra tính hợp lệ và kiểm tra email đã tồn tại hay chưa.",
            "Mật khẩu được băm trước khi lưu vào bảng users.",
            "Server tạo JWT và trả về thông tin tài khoản; người dùng được xác định là hội viên.",
            "Ở những lần tiếp theo, người dùng đăng nhập bằng email và mật khẩu để nhận token truy cập.",
        ],
    )
    add_numbered_process(
        doc,
        "b) Quy trình tìm hiểu và chọn gói tập",
        [
            "Khách hàng xem danh sách các gói tập đang hoạt động.",
            "Hệ thống hiển thị tên gói, giá, thời hạn, mô tả và các quyền lợi đi kèm.",
            "Khách hàng có thể lọc hoặc so sánh để lựa chọn gói phù hợp.",
            "Khi nhấn thêm vào giỏ hàng, hệ thống yêu cầu đăng nhập nếu khách hàng chưa có phiên hợp lệ.",
            "Gói tập được thêm vào giỏ; nếu gói đã tồn tại thì hệ thống tăng số lượng.",
        ],
    )
    add_numbered_process(
        doc,
        "c) Quy trình thanh toán và kích hoạt membership",
        [
            "Hội viên mở trang checkout; hệ thống lấy các gói trong giỏ và tính tạm tính.",
            "Hội viên lựa chọn phương thức thanh toán và có thể nhập mã khuyến mãi.",
            "Server kiểm tra thời gian hiệu lực, số lượt sử dụng, giá trị đơn tối thiểu và cách tính giảm giá của mã.",
            "Khi hội viên xác nhận, server tạo order và các order_items trong một giao dịch cơ sở dữ liệu.",
            "Hệ thống tạo membership tương ứng với từng gói, tính ngày hết hạn dựa trên duration_days và tạo thông báo cho hội viên.",
            "Số lần sử dụng khuyến mãi được cập nhật và giỏ hàng của hội viên được xóa.",
        ],
    )
    add_numbered_process(
        doc,
        "d) Quy trình theo dõi dịch vụ của hội viên",
        [
            "Hội viên truy cập dashboard sau khi đăng nhập.",
            "Hệ thống tải đồng thời danh sách membership, đơn hàng và thông báo của tài khoản.",
            "Hội viên xem trạng thái membership, thời hạn sử dụng, tổng số đơn hàng và thông báo chưa đọc.",
            "Dữ liệu hiển thị chỉ thuộc về tài khoản đang được xác thực.",
        ],
    )
    add_numbered_process(
        doc,
        "e) Quy trình quản trị dữ liệu",
        [
            "Quản trị viên đăng nhập bằng tài khoản admin và truy cập khu vực /admin.",
            "Middleware xác thực token và kiểm tra quyền admin trước khi xử lý yêu cầu.",
            "Quản trị viên chọn nhóm dữ liệu cần quản lý: hội viên, gói tập, huấn luyện viên, thiết bị, bảo trì, khuyến mãi hoặc đơn hàng.",
            "Hệ thống hiển thị danh sách; quản trị viên có thể thêm mới, cập nhật, thay đổi trạng thái hoặc xóa theo quy tắc nghiệp vụ.",
            "Sau mỗi thao tác thành công, danh sách được tải lại để phản ánh dữ liệu mới nhất.",
        ],
    )
    add_numbered_process(
        doc,
        "f) Quy trình quản lý bảo trì thiết bị",
        [
            "Quản trị viên cập nhật danh mục thiết bị với tên, loại, thương hiệu, số serial, tình trạng và vị trí.",
            "Khi thiết bị cần kiểm tra hoặc sửa chữa, quản trị viên tạo phiếu bảo trì và chọn thiết bị liên quan.",
            "Phiếu ghi nhận ngày bảo trì, loại công việc, mô tả, chi phí, nhân viên phụ trách và trạng thái.",
            "Sau khi hoàn thành, trạng thái phiếu được cập nhật để phục vụ theo dõi lịch sử bảo trì.",
        ],
    )

    doc.add_heading("1.3.4. Các biểu mẫu thu thập trong quá trình khảo sát", level=2)
    add_body_paragraph(
        doc,
        "Do IronZone là hệ thống mô phỏng phục vụ môn học, các biểu mẫu khảo sát được tổng hợp từ dữ liệu cần nhập trên giao diện và cấu trúc nghiệp vụ của website. Các biểu mẫu này là cơ sở để xác định thuộc tính dữ liệu, ràng buộc kiểm tra và thiết kế bảng trong cơ sở dữ liệu."
    )

    add_form_table(
        doc,
        "Bảng 1.1. Biểu mẫu đăng ký hội viên",
        [
            ("Họ và tên", "Tên đầy đủ của hội viên."),
            ("Email", "Địa chỉ đăng nhập, không được trùng với tài khoản đã tồn tại."),
            ("Mật khẩu", "Được kiểm tra độ dài và băm trước khi lưu."),
            ("Số điện thoại", "Thông tin liên hệ của hội viên."),
            ("Ngày sinh", "Dùng để hoàn thiện hồ sơ cá nhân."),
            ("Giới tính", "Nam, nữ hoặc khác."),
        ],
    )
    add_form_table(
        doc,
        "Bảng 1.2. Biểu mẫu đơn đăng ký gói tập",
        [
            ("Mã hội viên", "Xác định người tạo đơn."),
            ("Danh sách gói", "Gồm mã gói, tên gói, số lượng và đơn giá."),
            ("Mã khuyến mãi", "Không bắt buộc; được kiểm tra trước khi áp dụng."),
            ("Phương thức thanh toán", "Tiền mặt, chuyển khoản, MoMo, ZaloPay hoặc VNPay."),
            ("Tạm tính / giảm giá / tổng tiền", "Các giá trị do hệ thống tính toán."),
            ("Ghi chú", "Yêu cầu bổ sung của hội viên nếu có."),
        ],
    )
    add_form_table(
        doc,
        "Bảng 1.3. Biểu mẫu quản lý gói tập",
        [
            ("Tên gói", "Tên thương mại của gói tập."),
            ("Thời hạn", "Số ngày sử dụng dịch vụ."),
            ("Giá", "Giá bán của gói tập."),
            ("Mô tả", "Thông tin giới thiệu ngắn."),
            ("Danh sách quyền lợi", "Các dịch vụ được bao gồm trong gói."),
            ("Trạng thái", "Đang mở bán hoặc ngừng hoạt động."),
            ("Gói nổi bật", "Xác định gói được ưu tiên hiển thị."),
        ],
    )
    add_form_table(
        doc,
        "Bảng 1.4. Biểu mẫu quản lý bảo trì thiết bị",
        [
            ("Thiết bị", "Thiết bị cần kiểm tra, sửa chữa hoặc thay thế."),
            ("Nhân viên phụ trách", "Nhân viên được giao thực hiện công việc."),
            ("Ngày bảo trì", "Ngày dự kiến hoặc ngày thực hiện."),
            ("Loại bảo trì", "Định kỳ, sửa chữa, thay thế hoặc kiểm tra."),
            ("Mô tả công việc", "Nội dung cần thực hiện."),
            ("Chi phí", "Chi phí phát sinh cho lần bảo trì."),
            ("Trạng thái", "Đã lên lịch, đang thực hiện hoặc hoàn thành."),
        ],
    )

    doc.add_heading("1.4. KẾT CHƯƠNG", level=1)
    add_body_paragraph(
        doc,
        "Chương 1 đã giới thiệu bối cảnh và lý do lựa chọn đề tài quản lý phòng gym IronZone, đồng thời xác định mục tiêu và phạm vi phù hợp với học phần Phân tích thiết kế hệ thống. Qua quá trình khảo sát, chương đã mô tả sơ lược cấu trúc hệ thống thông tin, các nhóm đối tượng tham gia, cơ cấu tổ chức và những quy trình nghiệp vụ trọng tâm như đăng ký tài khoản, chọn gói tập, thanh toán, kích hoạt membership, quản trị dữ liệu và bảo trì thiết bị."
    )
    add_body_paragraph(
        doc,
        "Các biểu mẫu nghiệp vụ được tổng hợp đã giúp xác định những thông tin cần thu thập và lưu trữ trong hệ thống. Đây là cơ sở để tiếp tục xác định yêu cầu chức năng, yêu cầu phi chức năng, tác nhân và các trường hợp sử dụng. Chương 2 sẽ thực hiện phân tích hệ thống chi tiết, trong đó tập trung xây dựng mô hình Use Case, đặc tả các chức năng và mô tả luồng hoạt động của những quy trình quan trọng."
    )

    core = doc.core_properties
    core.title = "Báo cáo IronZone Gym - Chương 1: Tổng quan"
    core.subject = "Môn Phân tích thiết kế hệ thống"
    core.author = "Nhóm thực hiện IronZone Gym"
    core.keywords = "IronZone Gym, phân tích thiết kế hệ thống, khảo sát hệ thống"
    doc.save(OUTPUT_PATH)
    return OUTPUT_PATH


if __name__ == "__main__":
    print(build_document())
