from pathlib import Path
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

sys.path.insert(0, str(Path(__file__).resolve().parent))
from create_chapter1_report import (  # noqa: E402
    BLUE,
    DARK_BLUE,
    GRAY,
    NAVY,
    add_body_paragraph,
    add_bullet,
    add_page_number,
    configure_styles,
    set_cell_shading,
    set_repeat_table_header,
    set_run_font,
    set_table_fixed_width,
)


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "reports"
OUTPUT_PATH = OUTPUT_DIR / "Bao_cao_IronZone_Chuong_4.docx"


def add_code_block(doc, text):
    table = doc.add_table(rows=1, cols=1)
    set_table_fixed_width(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F2F4F7")
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.05
    for index, line in enumerate(text.splitlines()):
        if index:
            paragraph.add_run().add_break()
        run = paragraph.add_run(line)
        set_run_font(run, name="Consolas", size=9.5, color=RGBColor(35, 39, 43))
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(0)


def add_figure_placeholder(doc, caption):
    table = doc.add_table(rows=1, cols=1)
    set_table_fixed_width(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F7F8FA")
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(28)
    paragraph.paragraph_format.space_after = Pt(28)
    run = paragraph.add_run("[CHÈN ẢNH GIAO DIỆN TẠI ĐÂY]")
    set_run_font(run, size=11, color=GRAY, bold=True)
    cp = doc.add_paragraph(style="Figure Caption")
    cp.add_run(caption)


def add_numbered_steps(doc, steps):
    for step in steps:
        paragraph = doc.add_paragraph(style="List Number")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = paragraph.add_run(step)
        set_run_font(run)


def add_tech_table(doc):
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    set_table_fixed_width(table, [2100, 2500, 4760])
    for index, text in enumerate(("Thành phần", "Công nghệ", "Chức năng")):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, "E8EEF5")
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(text)
        set_run_font(run, size=11, color=NAVY, bold=True)
    set_repeat_table_header(table.rows[0])

    rows = [
        ("Frontend", "React 18", "Xây dựng giao diện và các thành phần tương tác."),
        ("Công cụ frontend", "Vite", "Khởi chạy môi trường phát triển và đóng gói ứng dụng."),
        ("Giao diện", "Tailwind CSS", "Định dạng bố cục và thiết kế giao diện."),
        ("Điều hướng", "React Router", "Điều hướng giữa các trang của ứng dụng."),
        ("Gọi API", "Axios", "Gửi yêu cầu và nhận dữ liệu từ server."),
        ("Backend", "Node.js, Express", "Xử lý nghiệp vụ và cung cấp REST API."),
        ("Cơ sở dữ liệu", "SQLite", "Lưu trữ dữ liệu của hệ thống."),
        ("Truy cập dữ liệu", "better-sqlite3", "Thực thi các câu truy vấn SQLite."),
        ("Xác thực", "JSON Web Token", "Xác thực phiên đăng nhập của người dùng."),
        ("Bảo mật", "bcryptjs", "Băm và kiểm tra mật khẩu."),
        ("Kiểm tra dữ liệu", "express-validator", "Kiểm tra dữ liệu đầu vào của API."),
        ("Triển khai", "Vercel, Render", "Triển khai frontend và backend trong môi trường demo."),
    ]
    for row in rows:
        cells = table.add_row().cells
        for index, text in enumerate(row):
            paragraph = cells[index].paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(text)
            set_run_font(run, size=10.5)
    caption = doc.add_paragraph(style="Figure Caption")
    caption.add_run("Bảng 4.1. Công nghệ sử dụng trong hệ thống IronZone")


def build_document():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(1.15)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.4)
    section.footer_distance = Inches(0.4)
    configure_styles(doc)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.paragraph_format.space_after = Pt(0)
    run = header.add_run("IRONZONE GYM | CÀI ĐẶT HỆ THỐNG")
    set_run_font(run, size=9, color=GRAY, bold=True)
    add_page_number(section.footer.paragraphs[0])

    chapter = doc.add_paragraph()
    chapter.alignment = WD_ALIGN_PARAGRAPH.CENTER
    chapter.paragraph_format.space_before = Pt(8)
    chapter.paragraph_format.space_after = Pt(6)
    run = chapter.add_run("CHƯƠNG 4")
    set_run_font(run, size=16, color=BLUE, bold=True)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(22)
    run = title.add_run("CÀI ĐẶT HỆ THỐNG")
    set_run_font(run, size=24, color=NAVY, bold=True)

    doc.add_heading("4.1. GIỚI THIỆU", level=1)
    add_body_paragraph(
        doc,
        "Sau khi hoàn thành quá trình phân tích và thiết kế, hệ thống quản lý phòng gym IronZone được xây dựng dưới dạng một ứng dụng web. Hệ thống gồm giao diện người dùng, máy chủ xử lý nghiệp vụ và cơ sở dữ liệu.",
    )
    add_body_paragraph(
        doc,
        "Chương này trình bày môi trường phát triển, công nghệ sử dụng, kiến trúc cài đặt, cấu trúc mã nguồn, cách tổ chức cơ sở dữ liệu và quá trình xây dựng các chức năng chính của hệ thống.",
    )

    doc.add_heading("4.2. MÔI TRƯỜNG VÀ CÔNG NGHỆ CÀI ĐẶT", level=1)
    add_body_paragraph(
        doc,
        "Hệ thống được phát triển bằng JavaScript cho cả phía client và server. Môi trường cài đặt yêu cầu Node.js phiên bản 20 trở lên và npm phiên bản 9 trở lên. Các công nghệ chính được trình bày trong bảng sau:",
    )
    add_tech_table(doc)

    doc.add_heading("4.3. KIẾN TRÚC CÀI ĐẶT", level=1)
    add_body_paragraph(
        doc,
        "IronZone được xây dựng theo kiến trúc client-server gồm ba thành phần chính: client, server và cơ sở dữ liệu. Client chịu trách nhiệm hiển thị giao diện và tiếp nhận thao tác; server xử lý nghiệp vụ và cung cấp REST API; cơ sở dữ liệu lưu trữ thông tin của hệ thống.",
    )
    add_numbered_steps(
        doc,
        [
            "Người dùng thực hiện thao tác trên giao diện React.",
            "Client gửi yêu cầu HTTP đến REST API bằng Axios.",
            "Express tiếp nhận yêu cầu và chuyển đến route phù hợp.",
            "Middleware thực hiện xác thực, phân quyền và kiểm tra dữ liệu.",
            "Server truy vấn hoặc cập nhật cơ sở dữ liệu SQLite.",
            "Kết quả được trả về client dưới dạng JSON.",
            "Client cập nhật giao diện và hiển thị thông báo cho người dùng.",
        ],
    )

    doc.add_heading("4.4. CẤU TRÚC MÃ NGUỒN", level=1)
    add_body_paragraph(doc, "Mã nguồn được chia thành hai phần độc lập là frontend và backend:")
    add_code_block(
        doc,
        """IronZone-Gym/
|-- client/
|-- server/
|-- README.md
`-- render.yaml""",
    )

    doc.add_heading("4.4.1. Cấu trúc frontend", level=2)
    add_code_block(
        doc,
        """client/src/
|-- api/
|-- components/
|-- context/
|-- hooks/
|-- pages/
|   `-- admin/
|-- App.jsx
|-- index.css
`-- main.jsx""",
    )
    add_bullet(doc, "api: chứa các hàm gửi yêu cầu đến REST API.")
    add_bullet(doc, "components: chứa các thành phần giao diện có thể tái sử dụng.")
    add_bullet(doc, "context: quản lý trạng thái đăng nhập, giỏ hàng và ngôn ngữ.")
    add_bullet(doc, "hooks: cung cấp các hook truy cập context.")
    add_bullet(doc, "pages: chứa các trang dành cho khách hàng, hội viên và quản trị viên.")
    add_bullet(doc, "App.jsx: khai báo các đường dẫn và cơ chế bảo vệ route.")
    add_bullet(doc, "main.jsx: khởi tạo React và các context provider.")

    doc.add_heading("4.4.2. Cấu trúc backend", level=2)
    add_code_block(
        doc,
        """server/
|-- db/
|-- middleware/
|-- routes/
|   `-- admin/
|-- utils/
`-- index.js""",
    )
    add_bullet(doc, "db: chứa kết nối cơ sở dữ liệu, cấu trúc bảng và dữ liệu mẫu.")
    add_bullet(doc, "middleware: chứa chức năng xác thực, phân quyền và kiểm tra dữ liệu.")
    add_bullet(doc, "routes: chứa các REST API dành cho người dùng.")
    add_bullet(doc, "routes/admin: chứa các REST API quản trị.")
    add_bullet(doc, "utils: chứa các hàm hỗ trợ JWT và định dạng phản hồi.")
    add_bullet(doc, "index.js: khởi tạo Express và đăng ký các nhóm route.")

    doc.add_heading("4.5. CÀI ĐẶT CƠ SỞ DỮ LIỆU", level=1)
    add_body_paragraph(
        doc,
        "Hệ thống sử dụng SQLite và thư viện better-sqlite3. Khi server khởi động, hàm migrate() đọc file schema.sql và tạo các bảng chưa tồn tại. Các khóa ngoại được bật để bảo đảm tính toàn vẹn dữ liệu.",
    )
    add_body_paragraph(doc, "Các bảng dữ liệu chính gồm:")
    for item in [
        "users, admins và employees",
        "trainers và plans",
        "cart_items, promotions, orders và order_items",
        "memberships",
        "equipment và maintenance_logs",
        "reviews và notifications",
    ]:
        add_bullet(doc, item)
    add_body_paragraph(
        doc,
        "Dữ liệu thử nghiệm được tạo bằng chương trình seed. Chương trình này cung cấp tài khoản mẫu, gói tập, huấn luyện viên, thiết bị, khuyến mãi, đơn hàng và membership để phục vụ quá trình chạy thử.",
    )
    add_code_block(doc, "cd server\nnpm run seed")

    doc.add_heading("4.6. CÀI ĐẶT CÁC CHỨC NĂNG CHÍNH", level=1)
    doc.add_heading("4.6.1. Đăng ký và đăng nhập", level=2)
    add_body_paragraph(
        doc,
        "Người dùng nhập thông tin tại giao diện đăng ký hoặc đăng nhập. Client gửi dữ liệu đến nhóm API /api/auth. Server kiểm tra dữ liệu, băm hoặc so sánh mật khẩu bằng bcryptjs. Khi đăng nhập thành công, server tạo JWT chứa mã tài khoản và loại người dùng.",
    )
    add_body_paragraph(
        doc,
        "Token được client lưu lại và tự động gửi trong tiêu đề Authorization của các yêu cầu cần xác thực.",
    )
    add_code_block(doc, "Authorization: Bearer <token>")

    doc.add_heading("4.6.2. Quản lý gói tập", level=2)
    add_body_paragraph(
        doc,
        "Danh sách gói tập được lấy từ API GET /api/plans. Khách hàng có thể xem tên, giá, thời hạn, mô tả và quyền lợi. Quản trị viên có thể thêm, sửa hoặc xóa gói tập thông qua nhóm API /api/admin/plans.",
    )
    add_figure_placeholder(doc, "Hình 4.1. Giao diện danh sách gói tập")

    doc.add_heading("4.6.3. Quản lý giỏ hàng", level=2)
    add_body_paragraph(
        doc,
        "Sau khi đăng nhập, hội viên có thể thêm gói tập vào giỏ hàng, thay đổi số lượng hoặc xóa gói. Hệ thống lưu giỏ hàng theo từng người dùng và tự động tính tổng tiền.",
    )
    add_code_block(
        doc,
        """GET    /api/cart
POST   /api/cart
PUT    /api/cart/:id
DELETE /api/cart/:id
DELETE /api/cart""",
    )
    add_figure_placeholder(doc, "Hình 4.2. Giao diện giỏ hàng")

    doc.add_heading("4.6.4. Thanh toán và tạo membership", level=2)
    add_body_paragraph(
        doc,
        "Tại trang checkout, hội viên lựa chọn phương thức thanh toán và có thể nhập mã khuyến mãi. Khi tạo đơn hàng, server thực hiện các thao tác trong một transaction nhằm bảo đảm dữ liệu được cập nhật đồng bộ.",
    )
    add_numbered_steps(
        doc,
        [
            "Lấy các gói tập trong giỏ hàng.",
            "Tính tạm tính và số tiền giảm giá.",
            "Tạo đơn hàng và các chi tiết đơn hàng.",
            "Tạo membership tương ứng với từng gói tập.",
            "Tạo thông báo xác nhận cho hội viên.",
            "Cập nhật số lượt sử dụng mã khuyến mãi.",
            "Xóa các gói tập khỏi giỏ hàng.",
        ],
    )
    add_figure_placeholder(doc, "Hình 4.3. Giao diện thanh toán")

    doc.add_heading("4.6.5. Dashboard hội viên", level=2)
    add_body_paragraph(
        doc,
        "Dashboard hiển thị membership, thời gian sử dụng, lịch sử đơn hàng và thông báo của hội viên. Dữ liệu chỉ được truy xuất cho tài khoản đang được xác thực.",
    )
    add_code_block(
        doc,
        """GET /api/memberships/my
GET /api/orders/my
GET /api/notifications/my""",
    )
    add_figure_placeholder(doc, "Hình 4.4. Giao diện dashboard hội viên")

    doc.add_heading("4.6.6. Trang quản trị", level=2)
    add_body_paragraph(
        doc,
        "Trang quản trị chỉ cho phép tài khoản admin truy cập. Middleware kiểm tra token và quyền trước khi xử lý yêu cầu. Quản trị viên có thể theo dõi số liệu tổng quan và quản lý hội viên, gói tập, huấn luyện viên, thiết bị, bảo trì, khuyến mãi cùng đơn hàng.",
    )
    add_figure_placeholder(doc, "Hình 4.5. Giao diện dashboard quản trị")

    doc.add_heading("4.6.7. Quản lý thiết bị và bảo trì", level=2)
    add_body_paragraph(
        doc,
        "Quản trị viên có thể cập nhật tên, loại, thương hiệu, số serial, giá mua, vị trí và tình trạng thiết bị. Khi cần bảo trì, quản trị viên tạo phiếu ghi nhận thiết bị, nhân viên phụ trách, ngày thực hiện, loại công việc, nội dung, chi phí và trạng thái.",
    )
    add_figure_placeholder(doc, "Hình 4.6. Giao diện quản lý thiết bị và bảo trì")

    doc.add_heading("4.7. HƯỚNG DẪN CÀI ĐẶT VÀ CHẠY HỆ THỐNG", level=1)
    doc.add_heading("4.7.1. Cài đặt thư viện", level=2)
    add_code_block(doc, "cd server\nnpm install\n\ncd ../client\nnpm install")

    doc.add_heading("4.7.2. Cấu hình biến môi trường", level=2)
    add_body_paragraph(doc, "Tạo file .env trong thư mục server với nội dung:")
    add_code_block(
        doc,
        """PORT=5000
JWT_SECRET=chuoi_bi_mat_co_do_dai_toi_thieu_32_ky_tu
JWT_EXPIRES_IN=7d
DB_PATH=./db/gym.db
NODE_ENV=development
CLIENT_URL=http://localhost:5173""",
    )
    add_body_paragraph(doc, "Tạo file .env trong thư mục client với nội dung:")
    add_code_block(doc, "VITE_API_URL=http://localhost:5000/api")

    doc.add_heading("4.7.3. Khởi tạo và chạy hệ thống", level=2)
    add_body_paragraph(doc, "Khởi tạo dữ liệu mẫu:")
    add_code_block(doc, "cd server\nnpm run seed")
    add_body_paragraph(doc, "Chạy backend:")
    add_code_block(doc, "cd server\nnpm run dev")
    add_body_paragraph(doc, "Chạy frontend trong cửa sổ terminal khác:")
    add_code_block(doc, "cd client\nnpm run dev")
    add_body_paragraph(
        doc,
        "Sau khi khởi động thành công, frontend hoạt động tại http://localhost:5173 và backend hoạt động tại http://localhost:5000.",
    )

    doc.add_heading("4.7.4. Tài khoản thử nghiệm", level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    set_table_fixed_width(table, [2100, 4260, 3000])
    for index, text in enumerate(("Vai trò", "Email", "Mật khẩu")):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, "E8EEF5")
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(text)
        set_run_font(run, size=11, color=NAVY, bold=True)
    for row in [
        ("Quản trị viên", "admin@ironzone.vn", "admin123"),
        ("Hội viên", "user1@example.com", "password123"),
    ]:
        cells = table.add_row().cells
        for index, text in enumerate(row):
            paragraph = cells[index].paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(text)
            set_run_font(run, size=11)

    doc.add_heading("4.8. KẾT CHƯƠNG", level=1)
    add_body_paragraph(
        doc,
        "Chương 4 đã trình bày quá trình cài đặt hệ thống quản lý phòng gym IronZone, bao gồm môi trường phát triển, công nghệ, kiến trúc, cấu trúc mã nguồn, cơ sở dữ liệu và cách triển khai các chức năng chính.",
    )
    add_body_paragraph(
        doc,
        "Kết quả cài đặt cho thấy hệ thống đã hỗ trợ các nghiệp vụ cơ bản như quản lý tài khoản, gói tập, giỏ hàng, đơn hàng, membership, thiết bị, bảo trì và các chức năng quản trị. Những kết quả này là cơ sở để đánh giá ưu điểm, hạn chế và hướng phát triển của hệ thống trong Chương 5.",
    )

    core = doc.core_properties
    core.title = "Báo cáo IronZone Gym - Chương 4: Cài đặt hệ thống"
    core.subject = "Môn Phân tích thiết kế hệ thống"
    core.author = "Nhóm thực hiện IronZone Gym"
    core.keywords = "IronZone Gym, cài đặt hệ thống, React, Express, SQLite"
    doc.save(OUTPUT_PATH)
    return OUTPUT_PATH


if __name__ == "__main__":
    print(build_document())
