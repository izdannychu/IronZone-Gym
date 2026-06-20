from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "reports"
OUTPUT_PATH = OUTPUT_DIR / "Bao_cao_IronZone_Chuong_1.docx"
LOGO_PATH = ROOT / "client" / "public" / "assets" / "ironzone-logo.png"

NAVY = RGBColor(24, 55, 83)
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
GRAY = RGBColor(94, 103, 112)
LIGHT_GRAY = "F4F6F9"
WHITE = RGBColor(255, 255, 255)


def set_run_font(run, name="Times New Roman", size=13, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)
    set_run_font(run, size=10, color=GRAY)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=110, start=140, bottom=110, end=140):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_fixed_width(table, widths):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def add_body_paragraph(doc, text, bold_prefix=None):
    paragraph = doc.add_paragraph(style="Normal")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if bold_prefix and text.startswith(bold_prefix):
        prefix = paragraph.add_run(bold_prefix)
        set_run_font(prefix, bold=True)
        remainder = paragraph.add_run(text[len(bold_prefix):])
        set_run_font(remainder)
    else:
        run = paragraph.add_run(text)
        set_run_font(run)
    return paragraph


def add_bullet(doc, text):
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = paragraph.add_run(text)
    set_run_font(run)
    return paragraph


def add_callout(doc, label, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = 0
    set_table_fixed_width(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_GRAY)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    label_run = paragraph.add_run(f"{label}: ")
    set_run_font(label_run, size=12, color=DARK_BLUE, bold=True)
    text_run = paragraph.add_run(text)
    set_run_font(text_run, size=12)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(13)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.3
    normal.paragraph_format.first_line_indent = Inches(0.4)

    heading_tokens = {
        "Heading 1": (18, NAVY, 18, 10),
        "Heading 2": (15, BLUE, 14, 7),
        "Heading 3": (13, DARK_BLUE, 10, 5),
    }
    for style_name, (size, color, before, after) in heading_tokens.items():
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(13)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(5)
        style.paragraph_format.line_spacing = 1.2

    if "Figure Caption" not in styles:
        caption = styles.add_style("Figure Caption", WD_STYLE_TYPE.PARAGRAPH)
    else:
        caption = styles["Figure Caption"]
    caption.font.name = "Times New Roman"
    caption._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    caption._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    caption.font.size = Pt(11)
    caption.font.italic = True
    caption.font.color.rgb = GRAY
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(8)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_cover(doc):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(18)
    paragraph.paragraph_format.space_after = Pt(16)
    if LOGO_PATH.exists():
        paragraph.add_run().add_picture(str(LOGO_PATH), width=Inches(2.2))

    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_before = Pt(24)
    kicker.paragraph_format.space_after = Pt(10)
    run = kicker.add_run("BÁO CÁO MÔN HỌC")
    set_run_font(run, size=14, color=BLUE, bold=True)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(12)
    run = title.add_run("PHÂN TÍCH VÀ THIẾT KẾ\nHỆ THỐNG IRONZONE GYM")
    set_run_font(run, size=25, color=NAVY, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(30)
    run = subtitle.add_run("Website quản lý phòng gym phục vụ minh họa các mô hình và sơ đồ phân tích thiết kế hệ thống")
    set_run_font(run, size=14, color=GRAY, italic=True)

    table = doc.add_table(rows=3, cols=2)
    set_table_fixed_width(table, [2700, 6660])
    labels = ["Học phần", "Sản phẩm minh họa", "Nội dung tài liệu"]
    values = [
        "Phân tích thiết kế hệ thống",
        "IronZone Gym Management System",
        "Chương 1 - Tổng quan",
    ]
    for row, label, value in zip(table.rows, labels, values):
        left, right = row.cells
        set_cell_shading(left, "E8EEF5")
        left_p = left.paragraphs[0]
        left_p.paragraph_format.space_after = Pt(0)
        left_run = left_p.add_run(label)
        set_run_font(left_run, size=12, color=NAVY, bold=True)
        right_p = right.paragraphs[0]
        right_p.paragraph_format.space_after = Pt(0)
        right_run = right_p.add_run(value)
        set_run_font(right_run, size=12)

    date = doc.add_paragraph()
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date.paragraph_format.space_before = Pt(55)
    date.paragraph_format.space_after = Pt(0)
    run = date.add_run("Năm 2026")
    set_run_font(run, size=12, color=GRAY)
    doc.add_page_break()


def build_document():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(1.15)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.4)
    section.footer_distance = Inches(0.4)

    configure_styles(doc)

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hp.paragraph_format.space_after = Pt(0)
    hr = hp.add_run("IRONZONE GYM | PHÂN TÍCH THIẾT KẾ HỆ THỐNG")
    set_run_font(hr, size=9, color=GRAY, bold=True)
    add_page_number(section.footer.paragraphs[0])

    add_cover(doc)

    chapter = doc.add_paragraph()
    chapter.alignment = WD_ALIGN_PARAGRAPH.CENTER
    chapter.paragraph_format.space_before = Pt(6)
    chapter.paragraph_format.space_after = Pt(6)
    run = chapter.add_run("CHƯƠNG 1")
    set_run_font(run, size=16, color=BLUE, bold=True)

    chapter_title = doc.add_paragraph()
    chapter_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    chapter_title.paragraph_format.space_after = Pt(22)
    run = chapter_title.add_run("TỔNG QUAN")
    set_run_font(run, size=24, color=NAVY, bold=True)

    add_callout(
        doc,
        "Định hướng chương",
        "Trình bày bối cảnh, lý do lựa chọn đề tài, bài toán đặt ra, mục tiêu, phạm vi và ý nghĩa của hệ thống IronZone Gym trong học phần Phân tích thiết kế hệ thống.",
    )

    doc.add_heading("1.1. Giới thiệu đề tài", level=2)
    add_body_paragraph(
        doc,
        "Trong quá trình chuyển đổi số, các cơ sở kinh doanh dịch vụ thể thao ngày càng có nhu cầu quản lý thông tin tập trung, giảm thao tác thủ công và cải thiện trải nghiệm khách hàng. Đối với một phòng gym, dữ liệu cần quản lý không chỉ bao gồm thông tin hội viên mà còn liên quan đến gói tập, huấn luyện viên, đơn hàng, khuyến mãi, thiết bị, lịch bảo trì và các thông báo dành cho khách hàng. Nếu các nội dung này được xử lý bằng sổ sách hoặc nhiều công cụ rời rạc, việc tra cứu, cập nhật và kiểm soát dữ liệu sẽ mất nhiều thời gian, đồng thời dễ phát sinh sai sót."
    )
    add_body_paragraph(
        doc,
        "Từ nhu cầu trên, nhóm lựa chọn xây dựng đề tài “Phân tích và thiết kế hệ thống quản lý phòng gym IronZone”. Sản phẩm của đề tài là một website full-stack mô phỏng các nghiệp vụ cơ bản của một phòng gym hiện đại. Hệ thống cung cấp giao diện cho khách hàng tìm hiểu gói tập, xem thông tin huấn luyện viên, đăng ký tài khoản, thêm gói tập vào giỏ hàng, áp dụng mã khuyến mãi, tạo đơn hàng và theo dõi membership. Bên cạnh đó, hệ thống còn có khu vực quản trị để quản lý hội viên, gói tập, huấn luyện viên, thiết bị, bảo trì, khuyến mãi và đơn hàng."
    )
    add_body_paragraph(
        doc,
        "Trong phạm vi học phần Phân tích thiết kế hệ thống, website IronZone không chỉ được xem là một sản phẩm phần mềm độc lập mà còn là tình huống nghiệp vụ mẫu. Các chức năng của website tạo cơ sở thực tế để xác định tác nhân, yêu cầu chức năng, yêu cầu phi chức năng, luồng xử lý, dữ liệu và mối quan hệ giữa các thành phần; từ đó xây dựng các sơ đồ phân tích và thiết kế phù hợp."
    )

    doc.add_heading("1.2. Lý do chọn đề tài", level=2)
    add_body_paragraph(
        doc,
        "Quản lý phòng gym là bài toán có phạm vi vừa đủ rộng để thể hiện đầy đủ các bước của quy trình phân tích và thiết kế hệ thống. Bài toán có nhiều nhóm người dùng, nhiều đối tượng dữ liệu và nhiều quy trình có sự liên kết chặt chẽ. Điều này giúp việc vận dụng kiến thức môn học trở nên trực quan hơn so với một ví dụ chỉ có ít chức năng hoặc thiếu luồng nghiệp vụ thực tế."
    )
    add_body_paragraph(doc, "Đề tài được lựa chọn dựa trên các lý do chính sau:")
    add_bullet(doc, "Bài toán gần gũi với thực tế và dễ khảo sát các nhu cầu cơ bản của khách hàng cũng như người quản trị phòng gym.")
    add_bullet(doc, "Hệ thống có hai nhóm tác nhân chính là hội viên và quản trị viên, thuận lợi cho việc xác định quyền hạn và xây dựng sơ đồ Use Case.")
    add_bullet(doc, "Các quy trình đăng ký, đăng nhập, mua gói tập, áp dụng khuyến mãi, tạo đơn hàng và kích hoạt membership phù hợp để mô hình hóa bằng Activity Diagram và Sequence Diagram.")
    add_bullet(doc, "Các đối tượng như User, Admin, Plan, Trainer, Order, Membership, Promotion và Equipment có quan hệ rõ ràng, phù hợp để xây dựng Class Diagram và mô hình dữ liệu ERD.")
    add_bullet(doc, "Website có cả phía người dùng, phía quản trị, máy chủ API và cơ sở dữ liệu, nhờ đó có thể minh họa kiến trúc hệ thống, sơ đồ thành phần và sơ đồ triển khai.")

    doc.add_heading("1.3. Bài toán đặt ra", level=2)
    add_body_paragraph(
        doc,
        "Một phòng gym cần cung cấp thông tin dịch vụ rõ ràng cho khách hàng, đồng thời duy trì khả năng quản lý dữ liệu vận hành. Khách hàng cần biết phòng gym đang có những gói tập nào, mức giá và quyền lợi ra sao, huấn luyện viên có chuyên môn gì, cách đăng ký và thanh toán như thế nào. Sau khi mua gói, khách hàng cần theo dõi trạng thái đơn hàng, thời hạn membership và các thông báo liên quan."
    )
    add_body_paragraph(
        doc,
        "Ở phía quản trị, hệ thống phải hỗ trợ cập nhật dữ liệu nghiệp vụ thường xuyên. Quản trị viên cần theo dõi số lượng thành viên, doanh thu, đơn hàng mới, thiết bị cần chú ý và các membership đang hoạt động. Đồng thời, quản trị viên phải có khả năng thêm, sửa, xóa hoặc thay đổi trạng thái của các đối tượng quản lý theo đúng quyền hạn."
    )
    add_body_paragraph(doc, "Từ đó, bài toán tổng quát được xác định như sau:")
    add_bullet(doc, "Xây dựng một hệ thống web tập trung, hỗ trợ khách hàng tiếp cận và đăng ký dịch vụ của phòng gym.")
    add_bullet(doc, "Tổ chức dữ liệu nhất quán để phục vụ các nghiệp vụ bán gói tập, quản lý membership và vận hành phòng gym.")
    add_bullet(doc, "Phân tách rõ quyền truy cập giữa người dùng thông thường và quản trị viên.")
    add_bullet(doc, "Đảm bảo các luồng nghiệp vụ có thể được mô tả, kiểm tra và đối chiếu thông qua các mô hình phân tích thiết kế hệ thống.")

    doc.add_heading("1.4. Mục tiêu của đề tài", level=2)
    doc.add_heading("1.4.1. Mục tiêu tổng quát", level=3)
    add_body_paragraph(
        doc,
        "Mục tiêu tổng quát của đề tài là phân tích, thiết kế và cài đặt thử nghiệm một hệ thống quản lý phòng gym trên nền tảng web. Hệ thống phải thể hiện được các nghiệp vụ tiêu biểu của IronZone, đồng thời đóng vai trò sản phẩm minh họa để đối chiếu giữa mô hình lý thuyết và chương trình thực tế trong học phần."
    )

    doc.add_heading("1.4.2. Mục tiêu cụ thể", level=3)
    add_bullet(doc, "Khảo sát và mô tả các nghiệp vụ chính của một hệ thống quản lý phòng gym.")
    add_bullet(doc, "Xác định tác nhân, yêu cầu chức năng, yêu cầu phi chức năng và các quy tắc nghiệp vụ.")
    add_bullet(doc, "Xây dựng các sơ đồ Use Case, Activity, Sequence, Class, ERD và các sơ đồ kiến trúc có liên quan.")
    add_bullet(doc, "Thiết kế cơ sở dữ liệu có khả năng biểu diễn đầy đủ các đối tượng và mối quan hệ nghiệp vụ.")
    add_bullet(doc, "Cài đặt website gồm giao diện khách hàng, giao diện quản trị, REST API và cơ sở dữ liệu.")
    add_bullet(doc, "Kiểm tra mức độ phù hợp giữa tài liệu phân tích thiết kế và chức năng đã được cài đặt.")

    doc.add_heading("1.5. Phạm vi đề tài", level=2)
    doc.add_heading("1.5.1. Phạm vi chức năng", level=3)
    add_body_paragraph(doc, "Trong phiên bản phục vụ báo cáo môn học, hệ thống tập trung vào các nhóm chức năng sau:")

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    set_table_fixed_width(table, [1850, 3160, 4350])
    headers = ["Nhóm chức năng", "Đối tượng sử dụng", "Nội dung chính"]
    for index, text in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, "E8EEF5")
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(text)
        set_run_font(run, size=11, color=NAVY, bold=True)
    set_repeat_table_header(table.rows[0])

    rows = [
        ("Thông tin công khai", "Khách truy cập", "Xem trang chủ, gói tập, huấn luyện viên, giới thiệu và thông tin liên hệ."),
        ("Tài khoản", "Hội viên, quản trị viên", "Đăng ký, đăng nhập, xác thực JWT, xem và cập nhật thông tin cá nhân."),
        ("Mua gói tập", "Hội viên", "Quản lý giỏ hàng, áp dụng mã khuyến mãi, chọn phương thức thanh toán và tạo đơn hàng."),
        ("Theo dõi dịch vụ", "Hội viên", "Xem membership, lịch sử đơn hàng và thông báo."),
        ("Quản trị vận hành", "Quản trị viên", "Quản lý hội viên, gói tập, HLV, thiết bị, bảo trì, khuyến mãi và đơn hàng."),
        ("Thống kê", "Quản trị viên", "Theo dõi thành viên, membership, doanh thu, trạng thái đơn hàng và thiết bị cần chú ý."),
    ]
    for group, actor, content in rows:
        cells = table.add_row().cells
        for index, text in enumerate((group, actor, content)):
            paragraph = cells[index].paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if index != 1 else WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(text)
            set_run_font(run, size=11)

    caption = doc.add_paragraph(style="Figure Caption")
    caption.add_run("Bảng 1.1. Phạm vi chức năng của hệ thống IronZone Gym")

    doc.add_heading("1.5.2. Phạm vi kỹ thuật", level=3)
    add_body_paragraph(
        doc,
        "Hệ thống được triển khai dưới dạng ứng dụng web full-stack. Phía client là Single Page Application xây dựng bằng React; phía server cung cấp REST API bằng Node.js và Express; dữ liệu được lưu trữ trong SQLite. Client và server giao tiếp qua HTTP/JSON. Việc xác thực sử dụng JSON Web Token, còn phân quyền được thực hiện theo loại tài khoản người dùng hoặc quản trị viên."
    )

    doc.add_heading("1.5.3. Giới hạn của đề tài", level=3)
    add_bullet(doc, "Hệ thống chủ yếu phục vụ mục đích học tập và trình diễn, chưa hướng đến vận hành thương mại quy mô lớn.")
    add_bullet(doc, "Các phương thức thanh toán được mô phỏng ở mức lựa chọn và ghi nhận trạng thái, chưa tích hợp cổng thanh toán thực tế.")
    add_bullet(doc, "Chưa xây dựng chức năng đặt lịch huấn luyện viên, check-in bằng thiết bị phần cứng hoặc quản lý lớp học theo thời khóa biểu.")
    add_bullet(doc, "SQLite phù hợp cho môi trường demo nhưng chưa phải lựa chọn tối ưu cho hệ thống có nhiều người dùng đồng thời.")
    add_bullet(doc, "Dữ liệu seed được sử dụng để minh họa nghiệp vụ, không đại diện cho dữ liệu vận hành thực tế.")

    doc.add_heading("1.6. Đối tượng sử dụng và các bên liên quan", level=2)
    add_body_paragraph(
        doc,
        "Việc xác định đúng các bên liên quan là cơ sở để thu thập yêu cầu và xây dựng mô hình Use Case. Trong hệ thống IronZone, các nhóm đối tượng chính được xác định như sau:"
    )
    add_bullet(doc, "Khách truy cập: người chưa đăng nhập, có nhu cầu tìm hiểu thông tin về phòng gym, gói tập và huấn luyện viên.")
    add_bullet(doc, "Hội viên: người đã có tài khoản, có thể mua gói tập, quản lý giỏ hàng, xem đơn hàng, membership và thông báo.")
    add_bullet(doc, "Quản trị viên: người vận hành hệ thống, có quyền truy cập trang quản trị và quản lý các dữ liệu nghiệp vụ.")
    add_bullet(doc, "Nhân viên phòng gym: đối tượng liên quan đến hoạt động vận hành và bảo trì thiết bị; dữ liệu nhân viên hiện được sử dụng trong nghiệp vụ bảo trì.")
    add_bullet(doc, "Nhóm phát triển và giảng viên: sử dụng hệ thống cùng tài liệu để đánh giá tính nhất quán giữa yêu cầu, mô hình thiết kế và sản phẩm cài đặt.")

    doc.add_heading("1.7. Phương pháp thực hiện", level=2)
    add_body_paragraph(
        doc,
        "Đề tài được thực hiện theo hướng tiếp cận phân tích hướng đối tượng, kết hợp giữa khảo sát nghiệp vụ, mô hình hóa và cài đặt thử nghiệm. Quy trình thực hiện gồm các bước:"
    )
    steps = [
        ("Khảo sát bài toán", "Xác định nhu cầu của khách hàng và quản trị viên trong hoạt động đăng ký, mua gói và vận hành phòng gym."),
        ("Thu thập và đặc tả yêu cầu", "Phân loại yêu cầu chức năng, yêu cầu phi chức năng, tác nhân và các quy tắc nghiệp vụ."),
        ("Phân tích hệ thống", "Xây dựng Use Case Diagram, đặc tả Use Case và Activity Diagram cho các quy trình quan trọng."),
        ("Thiết kế hệ thống", "Thiết kế kiến trúc, Sequence Diagram, Class Diagram, ERD, cơ sở dữ liệu và giao diện."),
        ("Cài đặt và kiểm tra", "Xây dựng client, server, cơ sở dữ liệu; chạy thử các luồng nghiệp vụ và đối chiếu với thiết kế."),
    ]
    for index, (name, description) in enumerate(steps, start=1):
        paragraph = doc.add_paragraph(style="List Number")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        name_run = paragraph.add_run(f"{name}: ")
        set_run_font(name_run, bold=True)
        desc_run = paragraph.add_run(description)
        set_run_font(desc_run)

    doc.add_heading("1.8. Công nghệ sử dụng", level=2)
    add_body_paragraph(
        doc,
        "Các công nghệ được lựa chọn dựa trên tiêu chí dễ triển khai, phù hợp với quy mô đồ án và thể hiện rõ sự phân chia giữa các tầng của hệ thống."
    )

    tech_table = doc.add_table(rows=1, cols=3)
    tech_table.style = "Table Grid"
    set_table_fixed_width(tech_table, [1900, 2700, 4760])
    for index, text in enumerate(("Thành phần", "Công nghệ", "Vai trò")):
        cell = tech_table.rows[0].cells[index]
        set_cell_shading(cell, "E8EEF5")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(text)
        set_run_font(r, size=11, color=NAVY, bold=True)
    set_repeat_table_header(tech_table.rows[0])
    tech_rows = [
        ("Frontend", "React, Vite, Tailwind CSS", "Xây dựng giao diện SPA, điều hướng trang và hiển thị dữ liệu."),
        ("Giao tiếp API", "Axios, HTTP/JSON", "Gửi yêu cầu từ client đến REST API và nhận phản hồi."),
        ("Backend", "Node.js, Express", "Xử lý nghiệp vụ, xác thực, phân quyền và cung cấp endpoint."),
        ("Cơ sở dữ liệu", "SQLite, better-sqlite3", "Lưu trữ dữ liệu và thực thi truy vấn có quan hệ."),
        ("Bảo mật", "bcryptjs, JWT", "Băm mật khẩu và xác thực phiên đăng nhập bằng token."),
        ("Kiểm tra dữ liệu", "express-validator", "Kiểm tra dữ liệu đầu vào trước khi xử lý nghiệp vụ."),
        ("Triển khai", "Vercel, Render", "Triển khai riêng frontend và backend trong môi trường demo."),
    ]
    for component, technology, role in tech_rows:
        cells = tech_table.add_row().cells
        for index, text in enumerate((component, technology, role)):
            p = cells[index].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(text)
            set_run_font(r, size=11)
    caption = doc.add_paragraph(style="Figure Caption")
    caption.add_run("Bảng 1.2. Các công nghệ chính được sử dụng")

    doc.add_heading("1.9. Ý nghĩa của đề tài", level=2)
    doc.add_heading("1.9.1. Ý nghĩa học thuật", level=3)
    add_body_paragraph(
        doc,
        "Đề tài giúp vận dụng một cách liên tục các nội dung của học phần, từ xác định vấn đề, phân tích yêu cầu đến thiết kế dữ liệu và kiến trúc. Thay vì xây dựng các sơ đồ độc lập, nhóm có thể đối chiếu từng sơ đồ với chức năng thực tế của website. Qua đó, mối liên hệ giữa Use Case, Activity, Sequence, Class Diagram, ERD và mã nguồn được thể hiện rõ ràng hơn."
    )
    doc.add_heading("1.9.2. Ý nghĩa thực tiễn", level=3)
    add_body_paragraph(
        doc,
        "Website IronZone mô phỏng một quy trình quản lý dịch vụ có thể bắt gặp trong thực tế. Sản phẩm giúp minh họa cách hệ thống thông tin hỗ trợ doanh nghiệp tập trung hóa dữ liệu, phân quyền người dùng, tự động hóa quy trình bán dịch vụ và cung cấp số liệu tổng quan cho người quản trị. Đây cũng là nền tảng để tiếp tục mở rộng các chức năng như đặt lịch PT, check-in, quản lý lớp học và tích hợp thanh toán."
    )

    doc.add_heading("1.10. Cấu trúc báo cáo", level=2)
    add_body_paragraph(doc, "Báo cáo được tổ chức thành năm chương, cụ thể như sau:")
    chapters = [
        ("Chương 1 - Tổng quan", "Giới thiệu đề tài, lý do lựa chọn, bài toán, mục tiêu, phạm vi, phương pháp thực hiện, công nghệ và ý nghĩa."),
        ("Chương 2 - Phân tích hệ thống", "Khảo sát nghiệp vụ, xác định yêu cầu, tác nhân, Use Case, đặc tả Use Case và các Activity Diagram."),
        ("Chương 3 - Thiết kế hệ thống", "Trình bày kiến trúc, thiết kế cơ sở dữ liệu, ERD, Class Diagram, Sequence Diagram, giao diện và các thành phần hệ thống."),
        ("Chương 4 - Cài đặt", "Mô tả môi trường phát triển, cấu trúc mã nguồn, quá trình cài đặt các chức năng và kết quả chạy thử."),
        ("Chương 5 - Kết luận", "Đánh giá kết quả đạt được, hạn chế của hệ thống và đề xuất hướng phát triển."),
    ]
    for name, description in chapters:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        name_run = paragraph.add_run(f"{name}: ")
        set_run_font(name_run, bold=True)
        description_run = paragraph.add_run(description)
        set_run_font(description_run)

    doc.add_heading("1.11. Tiểu kết chương", level=2)
    add_body_paragraph(
        doc,
        "Chương 1 đã trình bày bối cảnh và định hướng tổng thể của đề tài IronZone Gym. Hệ thống được lựa chọn vì có nghiệp vụ đủ đa dạng để phục vụ việc phân tích yêu cầu, mô hình hóa dữ liệu và thiết kế tương tác giữa các thành phần. Các mục tiêu, phạm vi và giới hạn đã được xác định nhằm bảo đảm đề tài phù hợp với thời lượng môn học. Trên cơ sở này, Chương 2 sẽ đi sâu vào phân tích hệ thống, xác định yêu cầu và mô hình hóa các chức năng chính."
    )

    core = doc.core_properties
    core.title = "Báo cáo IronZone Gym - Chương 1: Tổng quan"
    core.subject = "Môn Phân tích thiết kế hệ thống"
    core.author = "Nhóm thực hiện IronZone Gym"
    core.keywords = "IronZone Gym, phân tích thiết kế hệ thống, React, Express, SQLite"

    doc.save(OUTPUT_PATH)
    return OUTPUT_PATH


if __name__ == "__main__":
    print(build_document())
